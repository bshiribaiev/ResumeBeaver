import re
from fastapi import APIRouter, UploadFile, File, HTTPException, Response
from fastapi.responses import FileResponse
from pydantic import BaseModel
import os
import shutil
from datetime import datetime
import PyPDF2
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from resume_processor import ResumeAnalyzer
import tempfile
import uuid

router = APIRouter()

# Initialize analyzer (Watson will auto-detect availability)
analyzer = ResumeAnalyzer(use_watson=True)

# Request/Response models
class AnalyzeRequest(BaseModel):
    content: str
    type: str = "resume"  # "resume" or "job"

class OptimizeRequest(BaseModel):
    resume: str
    job_description: str

class MatchRequest(BaseModel):
    resume: str
    job_description: str

class GenerateResumeRequest(BaseModel):
    resume: str
    job_description: str
    applicant_name: str = "Applicant"
    format: str = "docx"  # "docx" or "txt"

# Extract text from uploaded file
def extract_file_text(file_path: str, file_type: str) -> str:
    try:
        if file_type == 'pdf':
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        elif file_type in ['docx', 'doc']:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text])
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells if cell.text])
                    if row_text:
                        text += "\n" + row_text
            return text.strip()
        else:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
    except Exception as e:
        raise Exception(f"Failed to extract text: {str(e)}")

def parse_resume_sections(resume_text: str) -> dict:
    """Parse resume text into structured sections"""
    sections = {}
    current_section = None
    current_content = []
    
    lines = resume_text.split('\n')
    
    # Common section headers
    section_keywords = {
        'SUMMARY': ['SUMMARY', 'OBJECTIVE', 'PROFILE', 'ABOUT'],
        'EXPERIENCE': ['EXPERIENCE', 'WORK EXPERIENCE', 'EMPLOYMENT', 'WORK HISTORY'],
        'EDUCATION': ['EDUCATION', 'ACADEMIC', 'LEARNING'],
        'SKILLS': ['SKILLS', 'TECHNICAL SKILLS', 'COMPETENCIES', 'TECHNOLOGIES'],
        'PROJECTS': ['PROJECTS', 'PROJECT EXPERIENCE'],
        'CERTIFICATIONS': ['CERTIFICATIONS', 'CERTIFICATES', 'AWARDS']
    }
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Check if this line is a section header
        is_section_header = False
        for section_name, keywords in section_keywords.items():
            if any(keyword in line.upper() for keyword in keywords):
                # Save previous section
                if current_section and current_content:
                    sections[current_section] = '\n'.join(current_content)
                
                # Start new section
                current_section = section_name
                current_content = []
                is_section_header = True
                break
        
        if not is_section_header and current_section:
            current_content.append(line)
    
    # Save the last section
    if current_section and current_content:
        sections[current_section] = '\n'.join(current_content)
    
    return sections

def format_experience_section(doc, experience_content: str):
    """Format experience section with proper job entry structure"""
    entries = experience_content.split('\n\n')  # Assume double newline separates jobs
    
    for entry in entries:
        if not entry.strip():
            continue
        
        lines = [line.strip() for line in entry.split('\n') if line.strip()]
        
        if lines:
            # First line should be job title/company
            job_para = doc.add_paragraph()
            job_run = job_para.add_run(lines[0])
            job_run.bold = True
            
            # Remaining lines are job details
            for line in lines[1:]:
                if line.startswith('•') or line.startswith('-'):
                    # Bullet point
                    doc.add_paragraph(line, style='List Bullet')
                else:
                    # Regular description
                    doc.add_paragraph(line)

def extract_contact_from_text(text: str) -> dict:
    """Extract contact information from resume text"""
    contact = {}
    
    # Email
    email_match = re.search(r'\b[\w._%+-]+@[\w.-]+\.[A-Z|a-z]{2,}\b', text)
    if email_match:
        contact['email'] = email_match.group(0)
    
    # Phone
    phone_match = re.search(r'(\+?1[-.\s]?)?\(?(\d{3})\)?[-.\s]?(\d{3})[-.\s]?(\d{4})', text)
    if phone_match:
        digits = ''.join(filter(str.isdigit, phone_match.group(0)))
        if len(digits) >= 10:
            contact['phone'] = f"({digits[-10:-7]}) {digits[-7:-4]}-{digits[-4:]}"
    
    # LinkedIn
    linkedin_match = re.search(r'linkedin\.com/in/([\w-]+)', text, re.IGNORECASE)
    if linkedin_match:
        contact['linkedin'] = f"linkedin.com/in/{linkedin_match.group(1)}"
    
    # GitHub
    github_match = re.search(r'github\.com/([\w-]+)', text, re.IGNORECASE)
    if github_match:
        contact['github'] = f"github.com/{github_match.group(1)}"
    
    return contact

def apply_optimizations_to_text(original_resume: str, optimization_result: dict) -> str:
    """Apply optimization suggestions to create improved resume text"""
    
    # Start with original resume
    optimized_text = original_resume
    
    # Get suggestions from optimization
    missing_skills = optimization_result.get('missing_skills', [])
    missing_keywords = optimization_result.get('missing_keywords', [])
    
    # Clean up and structure the resume better
    optimized_text = clean_resume_format(optimized_text)
    
    # Add missing skills to skills section
    if missing_skills:
        skills_to_add = missing_skills[:5]  # Add top 5 missing skills
        
        # Find or create skills section
        if re.search(r'\b(skills|technical skills|competencies)\b', optimized_text, re.IGNORECASE):
            # Add to existing skills section
            skills_pattern = r'(\b(?:SKILLS|TECHNICAL SKILLS|COMPETENCIES)\b.*?)(\n\n|\n[A-Z][A-Z]|\Z)'
            skills_match = re.search(skills_pattern, optimized_text, re.IGNORECASE | re.DOTALL)
            if skills_match:
                skills_section = skills_match.group(1)
                # Add skills in a clean format
                enhanced_skills = skills_section.rstrip() + f"\n• {' • '.join(skills_to_add)}"
                optimized_text = optimized_text.replace(skills_section, enhanced_skills)
        else:
            # Add new skills section before education/experience
            skills_section = f"\n\nSKILLS\n• {' • '.join(skills_to_add)}"
            # Insert before EDUCATION section if it exists
            if 'EDUCATION' in optimized_text.upper():
                optimized_text = re.sub(r'(\nEDUCATION)', skills_section + r'\1', optimized_text, flags=re.IGNORECASE)
            else:
                optimized_text += skills_section
    
    # Enhance summary with keywords
    if missing_keywords:
        keywords_to_add = missing_keywords[:3]  # Add top 3 keywords
        summary_pattern = r'(\b(?:SUMMARY|OBJECTIVE|PROFILE)\b.*?)(\n\n|\n[A-Z][A-Z]|\Z)'
        summary_match = re.search(summary_pattern, optimized_text, re.IGNORECASE | re.DOTALL)
        if summary_match:
            summary_section = summary_match.group(1)
            enhanced_summary = summary_section.rstrip() + f" Experienced with {', '.join(keywords_to_add)}."
            optimized_text = optimized_text.replace(summary_section, enhanced_summary)
    
    return optimized_text

def create_docx_resume(resume_text: str, applicant_name: str, optimization_data: dict) -> str:
    """Create a professional DOCX resume with proper formatting"""
    
    # Create document
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    # Apply optimizations to get improved content
    optimized_text = apply_optimizations_to_text(resume_text, optimization_data)
    
    # Parse the resume content to extract structured sections
    resume_sections = parse_resume_sections(optimized_text)
    
    # Header with name (larger, bold, centered)
    name_para = doc.add_heading(applicant_name.upper(), 0)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add contact information if available
    contact_info = extract_contact_from_text(optimized_text)
    if contact_info:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_parts = []
        if contact_info.get('email'):
            contact_parts.append(contact_info['email'])
        if contact_info.get('phone'):
            contact_parts.append(contact_info['phone'])
        if contact_info.get('linkedin'):
            contact_parts.append(contact_info['linkedin'])
        if contact_info.get('github'):
            contact_parts.append(contact_info['github'])
        
        contact_para.add_run(" • ".join(contact_parts))
        doc.add_paragraph()  # Empty line
    
    # Add optimization score
    if optimization_data.get('match_score'):
        match_score = optimization_data['match_score'].get('overall_score', 0)
        score_para = doc.add_paragraph()
        score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        score_run = score_para.add_run(f"Optimized for Target Position • Match Score: {match_score}%")
        score_run.italic = True
        doc.add_paragraph()  # Empty line
    
    # Add sections in proper order
    section_order = ['SUMMARY', 'EXPERIENCE', 'EDUCATION', 'SKILLS', 'PROJECTS', 'CERTIFICATIONS']
    
    for section_name in section_order:
        section_content = resume_sections.get(section_name)
        if section_content:
            # Section header
            doc.add_heading(section_name, level=1)
            
            # Section content
            if section_name == 'SKILLS':
                # Format skills nicely
                skills_para = doc.add_paragraph()
                skills_para.add_run(section_content)
            elif section_name == 'EXPERIENCE':
                # Format experience entries
                format_experience_section(doc, section_content)
            else:
                # Regular paragraph formatting
                for line in section_content.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())
    
    # Add improvements section
    if optimization_data.get('suggestions'):
        doc.add_page_break()
        doc.add_heading('OPTIMIZATION SUMMARY', level=1)
        
        suggestions = optimization_data.get('suggestions', [])
        for suggestion in suggestions[:3]:
            bullet_para = doc.add_paragraph()
            bullet_para.add_run(f"• {suggestion.get('category', '')}: ").bold = True
            bullet_para.add_run(suggestion.get('suggestion', ''))
    
    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    
    return temp_file.name

def create_txt_resume(resume_text: str, applicant_name: str, optimization_data: dict) -> str:
    """Create a plain text ATS-friendly resume"""
    
    output = []
    
    # Header
    output.append(applicant_name.upper())
    output.append("=" * len(applicant_name))
    output.append("")
    
    # Add optimization info
    if optimization_data.get('match_score'):
        match_score = optimization_data['match_score'].get('overall_score', 0)
        output.append(f"Resume optimized for target position (Match Score: {match_score}%)")
        output.append("")
    
    # Add optimized resume content
    output.append(resume_text)
    
    # Add suggestions
    if optimization_data.get('suggestions'):
        output.append("\n" + "="*50)
        output.append("OPTIMIZATION APPLIED:")
        output.append("="*50)
        
        suggestions = optimization_data.get('suggestions', [])
        for i, suggestion in enumerate(suggestions[:3], 1):
            output.append(f"{i}. {suggestion.get('category', '')}: {suggestion.get('suggestion', '')}")
    
    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.txt', mode='w', encoding='utf-8')
    temp_file.write('\n'.join(output))
    temp_file.close()
    
    return temp_file.name

def clean_resume_format(text: str) -> str:
    """Clean up resume formatting for better structure"""
    
    # Remove excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Ensure section headers are properly formatted
    section_headers = ['SUMMARY', 'OBJECTIVE', 'EXPERIENCE', 'EDUCATION', 'SKILLS', 'PROJECTS', 'CERTIFICATIONS']
    for header in section_headers:
        # Make sure section headers are on their own line and uppercase
        text = re.sub(rf'\b{header}\b', f'\n\n{header}', text, flags=re.IGNORECASE)
        text = re.sub(rf'\n\n{header.lower()}\b', f'\n\n{header}', text)
    
    # Clean up multiple spaces
    text = re.sub(r' {2,}', ' ', text)
    
    # Remove leading/trailing whitespace
    text = text.strip()
    
    return text

# Upload and process resume file
@router.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    allowed_extensions = ['.pdf', '.docx', '.doc', '.txt']
    file_ext = os.path.splitext(file.filename)[1].lower()
    
    if file_ext not in allowed_extensions:
        raise HTTPException(400, f"File type not supported. Allowed: {allowed_extensions}")
    if file.size > 10 * 1024 * 1024:  # 10MB limit
        raise HTTPException(400, "File too large (max 10MB)")
    try:
        # Save file temporarily
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        upload_dir = "uploads"
        os.makedirs(upload_dir, exist_ok=True)
        file_path = os.path.join(upload_dir, f"{timestamp}_{file.filename}")
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        # Extract and analyze text
        text = extract_file_text(file_path, file_ext[1:])
        analysis = analyzer.analyze_resume(text)
        return {
            "success": True,
            "filename": file.filename,
            "file_size": file.size,
            "text_preview": text[:500] + "..." if len(text) > 500 else text,
            "analysis": analysis,
            "timestamp": datetime.now().isoformat()
        }
    except Exception as e:
        # Clean up on error
        if 'file_path' in locals() and os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(500, f"Processing failed: {str(e)}")

# Analyze resume or job description
@router.post("/analyze")
async def analyze_content(request: AnalyzeRequest):
    try:
        if request.type == "resume":
            result = analyzer.analyze_resume(request.content)
        else:  # job description
            skills = analyzer.extract_skills(request.content)
            # Extract requirements
            exp_matches = re.findall(r'(\d+)\+?\s*years?', request.content, re.IGNORECASE)
            years_required = max(map(int, exp_matches)) if exp_matches else None
            result = {
                "skills_required": skills,
                "years_experience": years_required,
                "word_count": len(request.content.split()),
                "type": "job_description"
            }
        return {
            "success": True,
            "analysis": result,
            "ai_powered": result.get('ai_powered', False)
        }
    except Exception as e:
        raise HTTPException(500, f"Analysis failed: {str(e)}")

# Calculate match score between resume and job
@router.post("/match")
async def calculate_match(request: MatchRequest):
    try:
        result = analyzer.calculate_match_score(
            request.resume,
            request.job_description
        )
        return {
            "success": True,
            "match_analysis": result,
            "timestamp": datetime.now().isoformat()
        }
    except Exception as e:
        raise HTTPException(500, f"Match calculation failed: {str(e)}")

# Optimize resume for specific job
@router.post("/optimize")
async def optimize_resume(request: OptimizeRequest):
    try:
        result = analyzer.optimize_resume(
            request.resume,
            request.job_description
        )
        return {
            "success": True,
            "optimization": result,
            "ai_powered": result.get('ai_powered', False),
            "timestamp": datetime.now().isoformat()
        }
    except Exception as e:
        raise HTTPException(500, f"Optimization failed: {str(e)}")

# Generate optimized resume
@router.post("/generate")
async def generate_optimized_resume(request: GenerateResumeRequest):
    try:
        # Get optimization analysis
        optimization_result = analyzer.optimize_resume(
            request.resume,
            request.job_description
        )
        
        # Apply optimizations to create improved resume
        optimized_text = apply_optimizations_to_text(request.resume, optimization_result)
        
        # Generate file based on format
        if request.format.lower() == 'docx':
            file_path = create_docx_resume(optimized_text, request.applicant_name, optimization_result)
            media_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            filename = f"{request.applicant_name.replace(' ', '_')}_optimized_resume.docx"
        else:
            file_path = create_txt_resume(optimized_text, request.applicant_name, optimization_result)
            media_type = 'text/plain'
            filename = f"{request.applicant_name.replace(' ', '_')}_optimized_resume.txt"
        
        # Return file for download
        return FileResponse(
            path=file_path,
            media_type=media_type,
            filename=filename
        )
        
    except Exception as e:
        raise HTTPException(500, f"Resume generation failed: {str(e)}")

# Preview optimized resume
@router.post("/preview")
async def preview_optimized_resume(request: OptimizeRequest):
    try:
        # Get optimization analysis
        optimization_result = analyzer.optimize_resume(
            request.resume,
            request.job_description
        )
        
        # Apply optimizations to create improved resume
        optimized_text = apply_optimizations_to_text(request.resume, optimization_result)
        
        return {
            "success": True,
            "original_resume": request.resume,
            "optimized_resume": optimized_text,
            "optimization_data": optimization_result,
            "improvements": {
                "skills_added": optimization_result.get('missing_skills', [])[:5],
                "keywords_added": optimization_result.get('missing_keywords', [])[:8],
                "match_score_improvement": "Estimated +15-25% improvement"
            },
            "timestamp": datetime.now().isoformat()
        }
    except Exception as e:
        raise HTTPException(500, f"Preview generation failed: {str(e)}")

# Check system status
@router.get("/status")
async def get_status():
    watson_available = False
    model = "fallback_mode"
    
    try:
        from watson_client import get_watson_client
        client = get_watson_client()
        watson_available = client.watson_available
        
        # FIXED: Get actual model from environment variables
        if watson_available:
            model = os.getenv("IBM_MODEL_ID", "ibm/granite-13b-instruct-v2")
    except Exception as e:
        print(f"Watson client error: {e}")
    
    return {
        "status": "operational",
        "version": "2.0.0",
        "features": {
            "file_upload": True,
            "resume_analysis": True,
            "job_analysis": True,
            "match_scoring": True,
            "optimization": True,
            "resume_generation": True,
            "ai_powered": watson_available
        },
        "ai_status": {
            "watson_available": watson_available,
            "model": model,  # Now shows correct model: ibm/granite-13b-instruct-v2
            "model_provider": "IBM Watson" if watson_available else "Local fallback"
        },
        "endpoints": [
            "POST /upload - Upload and process resume file",
            "POST /analyze - Analyze resume or job content",
            "POST /match - Calculate resume-job match score",
            "POST /optimize - Generate optimization suggestions",
            "POST /generate - Generate optimized resume file",
            "POST /preview - Preview optimized resume",
            "GET /status - System status and capabilities"
        ]
    }

# Health check
@router.get("/health")
async def health_check():
    return {"status": "healthy", "service": "Resume Optimizer API"}