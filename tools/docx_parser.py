# tools/docx_parser.py
from docx import Document
import logging

def extract_docx_text(file_path: str) -> str:
    """
    Extract structured text from DOCX resume files.
    
    Enhanced DOCX processing that preserves formatting context
    and handles tables, headers, and complex layouts.
    
    Args:
        file_path (str): Path to DOCX file
        
    Returns:
        str: Extracted and structured text content
        
    Raises:
        Exception: If DOCX cannot be processed
    """
    text = ""
    try:
        doc = Document(file_path)
        
        # Enhanced: Process paragraphs with better formatting
        for para in doc.paragraphs:
            if para.text.strip():  # Skip empty paragraphs
                text += para.text + "\n"
        
        # Enhanced: Extract text from tables (common in resumes)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"
        
        # Enhanced: Log processing info
        logging.info(f"Processed DOCX: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
        
        return text.strip()
        
    except Exception as e:
        logging.error(f"DOCX extraction error for {file_path}: {str(e)}")
        raise Exception(f"Error reading DOCX: {str(e)}")

# Keep your original function name for backward compatibility
def extract_text_from_docx(file_path: str) -> str:
    """Backward compatibility wrapper"""
    return extract_docx_text(file_path)